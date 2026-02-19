"""
文档解析与分块服务 —— 将上传的文件拆分为适合检索的文本块

支持 5 种分块方式：
1. auto     - 自动段落（默认，按自然段落合并到 ~500 字）
2. heading1 - 按一级标题分章
3. heading2 - 按二级标题分段
4. qa       - 问答对（自动识别 Q/A 结构）
5. page     - 逐页分割（仅 PDF 有效，其他格式退化为自动段落）
"""

import os
import re
from typing import List, Dict
from utils import logger


def parse_txt(file_path: str) -> str:
    """解析纯文本 / Markdown 文件"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    # fallback
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def parse_docx(file_path: str) -> str:
    """解析 Word 文档（参考 swxy 的 Docx 类，简化版）"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                # 保留标题标记，便于按标题分块
                if p.style and p.style.name:
                    style = p.style.name.lower()
                    if 'heading 1' in style or style == 'heading 1':
                        paragraphs.append(f"# {text}")
                    elif 'heading 2' in style or style == 'heading 2':
                        paragraphs.append(f"## {text}")
                    elif 'heading 3' in style or style == 'heading 3':
                        paragraphs.append(f"### {text}")
                    else:
                        paragraphs.append(text)
                else:
                    paragraphs.append(text)
        # 也提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return '\n'.join(paragraphs)
    except Exception as e:
        logger.error(f"解析 DOCX 失败: {e}")
        return ""


def parse_pdf(file_path: str) -> str:
    """解析 PDF 文件"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return '\n'.join(pages)
    except Exception as e:
        logger.error(f"解析 PDF 失败: {e}")
        return ""


def parse_pdf_by_page(file_path: str) -> List[str]:
    """解析 PDF 文件，返回按页拆分的列表"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())
        return pages
    except Exception as e:
        logger.error(f"解析 PDF（逐页）失败: {e}")
        return []


def parse_document(file_path: str, file_type: str) -> str:
    """
    根据文件类型解析文档，返回纯文本内容。

    :param file_path: 文件路径
    :param file_type: 文件扩展名 (txt, md, docx, pdf)
    :return: 文档纯文本
    """
    file_type = file_type.lower()
    if file_type in ('txt', 'md', 'markdown'):
        return parse_txt(file_path)
    elif file_type in ('docx',):
        return parse_docx(file_path)
    elif file_type in ('pdf',):
        return parse_pdf(file_path)
    elif file_type in ('doc',):
        logger.warning("不支持 .doc 格式，请转换为 .docx")
        return ""
    else:
        logger.warning(f"不支持的文件格式: {file_type}")
        return ""


# ==================== 5 种分块策略 ====================

def split_auto(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[str]:
    """
    策略1: 自动段落分块（默认）
    按自然段落合并到 chunk_size 大小。
    """
    if not text or not text.strip():
        return []

    separators = ['\n\n', '\n', '。', '！', '？', '；', '.', '!', '?', ';']
    paragraphs = _split_text(text, separators)

    chunks = []
    current_chunk = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if len(para) > chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""
            for i in range(0, len(para), chunk_size - chunk_overlap):
                sub = para[i:i + chunk_size]
                if sub.strip():
                    chunks.append(sub.strip())
            continue

        if len(current_chunk) + len(para) + 1 > chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())
            if chunk_overlap > 0 and current_chunk:
                overlap_text = current_chunk[-chunk_overlap:]
                current_chunk = overlap_text + " " + para
            else:
                current_chunk = para
        else:
            current_chunk = current_chunk + "\n" + para if current_chunk else para

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return [c for c in chunks if len(c) >= 10]


def split_by_heading(text: str, level: int = 1) -> List[str]:
    """
    策略2/3: 按标题分块
    level=1 → 按一级标题(#)分章
    level=2 → 按二级标题(##)分段
    """
    if not text or not text.strip():
        return []

    if level == 1:
        # 匹配 # 开头但不是 ## 的行
        pattern = r'^# (?!#)'
    else:
        # 匹配 ## 开头但不是 ### 的行
        pattern = r'^#{1,2} (?!#)'

    lines = text.split('\n')
    chunks = []
    current_chunk_lines = []
    current_heading = ""

    for line in lines:
        is_heading = bool(re.match(pattern, line.strip(), re.MULTILINE))
        if is_heading:
            # 保存上一块
            if current_chunk_lines:
                chunk_text = '\n'.join(current_chunk_lines).strip()
                if chunk_text and len(chunk_text) >= 10:
                    chunks.append(chunk_text)
            current_chunk_lines = [line]
            current_heading = line.strip()
        else:
            current_chunk_lines.append(line)

    # 保存最后一块
    if current_chunk_lines:
        chunk_text = '\n'.join(current_chunk_lines).strip()
        if chunk_text and len(chunk_text) >= 10:
            chunks.append(chunk_text)

    # 如果没有找到标题，退化为自动段落
    if len(chunks) <= 1 and len(text) > 500:
        logger.info("未检测到标题结构，退化为自动段落分块")
        return split_auto(text)

    return chunks


def split_by_qa(text: str) -> List[str]:
    """
    策略4: 问答对分块
    自动识别多种 Q/A 结构，每对 QA 为一个 chunk。

    支持的格式：
    - ### 1、问题？  （数字编号标题，最常见）
    - ## 1. 问题
    - Q: ... A: ...
    - 问：... 答：...
    - 1、问题？  2、问题？
    - 【问题】...【回答思路】...
    """
    if not text or not text.strip():
        return []

    lines = text.split('\n')
    chunks = []

    # ---------- 策略 A：按带数字编号的标题行拆分 ----------
    # 匹配 ### 1、  ### 2.  ## 1、  # 1.  等
    heading_number_pattern = re.compile(
        r'^#{1,4}\s*(\d+)\s*[、.．)\]）】]\s*'
    )
    # 也匹配纯数字编号行：1、问题  1. 问题  (1) 问题
    plain_number_pattern = re.compile(
        r'^(\d+)\s*[、.．)\]）】]\s*\S'
    )

    # 先检测是否存在带数字的标题行
    heading_indices = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if heading_number_pattern.match(stripped) or plain_number_pattern.match(stripped):
            heading_indices.append(i)

    if len(heading_indices) >= 2:
        # 按标题行拆分，每个标题 + 其下内容为一个 chunk
        for idx, start in enumerate(heading_indices):
            end = heading_indices[idx + 1] if idx + 1 < len(heading_indices) else len(lines)
            chunk_text = '\n'.join(lines[start:end]).strip()
            if chunk_text and len(chunk_text) >= 10:
                chunks.append(chunk_text)

        if chunks:
            # 如果标题行之前有内容（如文档标题/说明），也保留
            if heading_indices[0] > 0:
                preamble = '\n'.join(lines[:heading_indices[0]]).strip()
                if preamble and len(preamble) >= 10:
                    chunks.insert(0, preamble)
            logger.info(f"QA 分块：按数字编号标题拆分，共 {len(chunks)} 个块")
            return chunks

    # ---------- 策略 B：按 Q/A 或 问/答 配对拆分 ----------
    qa_start_pattern = re.compile(
        r'^(?:Q|问|问题|Question)\s*[:：\d.、]*\s*',
        re.IGNORECASE
    )

    qa_indices = []
    for i, line in enumerate(lines):
        if qa_start_pattern.match(line.strip()):
            qa_indices.append(i)

    if len(qa_indices) >= 2:
        for idx, start in enumerate(qa_indices):
            end = qa_indices[idx + 1] if idx + 1 < len(qa_indices) else len(lines)
            chunk_text = '\n'.join(lines[start:end]).strip()
            if chunk_text and len(chunk_text) >= 10:
                chunks.append(chunk_text)

        if chunks:
            logger.info(f"QA 分块：按 Q/问 标记拆分，共 {len(chunks)} 个块")
            return chunks

    # ---------- 策略 C：退化为自动分块 ----------
    logger.info("未检测到 QA 结构，退化为自动段落分块")
    return split_auto(text)


def split_by_page(file_path: str, file_type: str) -> List[str]:
    """
    策略5: 逐页分割（主要适用于 PDF）
    非 PDF 格式退化为自动段落。
    """
    if file_type.lower() == 'pdf':
        pages = parse_pdf_by_page(file_path)
        if pages:
            return [p for p in pages if len(p) >= 10]

    # 非 PDF 或解析失败，退化为自动段落
    logger.info(f"文件类型 {file_type} 不支持逐页分割，退化为自动段落分块")
    text = parse_document(file_path, file_type)
    return split_auto(text)


# ==================== 辅助函数 ====================

def _split_text(text: str, separators: List[str]) -> List[str]:
    """按优先级选择分割符来拆分文本"""
    if not separators:
        return [text]

    sep = separators[0]
    remaining_seps = separators[1:]

    parts = text.split(sep)

    result = []
    for part in parts:
        if len(part) > 1000 and remaining_seps:
            result.extend(_split_text(part, remaining_seps))
        else:
            result.append(part)

    return result


# ==================== 主入口 ====================

# 分块方式映射表
CHUNK_METHODS = {
    "auto": "自动段落",
    "heading1": "按标题分章",
    "heading2": "按小节分段",
    "qa": "问答对",
    "page": "逐页分割",
}


def process_document(
    file_path: str,
    file_type: str,
    chunk_method: str = "auto",
) -> List[Dict]:
    """
    完整的文档处理流程：解析 → 分块。

    :param file_path: 文件路径
    :param file_type: 文件类型
    :param chunk_method: 分块方式 (auto/heading1/heading2/qa/page)
    :return: [{"chunk_index": 0, "content": "..."}]
    """
    method_name = CHUNK_METHODS.get(chunk_method, "自动段落")
    logger.info(f"文档分块方式: {method_name} ({chunk_method})")

    # 逐页分割需要直接操作文件
    if chunk_method == "page":
        raw_chunks = split_by_page(file_path, file_type)
    else:
        text = parse_document(file_path, file_type)
        if not text:
            return []

        if chunk_method == "heading1":
            raw_chunks = split_by_heading(text, level=1)
        elif chunk_method == "heading2":
            raw_chunks = split_by_heading(text, level=2)
        elif chunk_method == "qa":
            raw_chunks = split_by_qa(text)
        else:
            raw_chunks = split_auto(text)

    result = []
    for i, chunk in enumerate(raw_chunks):
        result.append({
            "chunk_index": i,
            "content": chunk,
        })

    logger.info(f"文档解析完成: {os.path.basename(file_path)}, 分块方式: {method_name}, 共 {len(result)} 个切片")
    return result
